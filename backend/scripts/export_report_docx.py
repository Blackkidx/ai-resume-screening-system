# -*- coding: utf-8 -*-
"""
📄 Export Academic Report to DOCX
Generates Word document with embedded figures for thesis.
"""
import sys
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT

from pathlib import Path

REPORT_DIR = Path("d:/ai-resume-screening-system/reports")
OUTPUT_PATH = REPORT_DIR / "XGBoost_v5_Training_Report.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def add_image(doc, filename, caption, width=Inches(5.5)):
    img_path = REPORT_DIR / filename
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        add_para(doc, f"[Image not found: {filename}]", italic=True)


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "TH SarabunPSK"
    font.size = Pt(14)

    # ═══════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("รายงานผลการเทรนโมเดล XGBoost v5.0")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI Resume Screening System")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Model: XGBoost (Extreme Gradient Boosting)\n"
                       "Version: v5.0 (Final)\n"
                       "Training Date: 2026-03-05\n"
                       "Data: 805 records (585 real + 220 synthetic)")
    run.font.size = Pt(14)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════
    add_heading(doc, "1. บทนำ (Introduction)")
    add_para(doc, "ระบบคัดกรอง Resume อัตโนมัติใช้ XGBoost (Extreme Gradient Boosting) เป็น Machine Learning model สำหรับจำแนกใบสมัครฝึกงานเป็น 2 กลุ่ม: ผ่าน (Accepted) และ ไม่ผ่าน (Rejected) โดยเรียนรู้จากข้อมูลการตัดสินใจจริงของ HR ผู้เชี่ยวชาญ 10 ตำแหน่ง 5 บริษัท")
    add_para(doc, "XGBoost เป็น Ensemble Learning แบบ Gradient Boosted Decision Trees ที่มีประสิทธิภาพสูงสำหรับ Tabular Data Classification ถูกเลือกใช้เนื่องจากมีความสามารถในการจัดการ features ที่หลากหลาย, ทนต่อ missing values, และมีกลไก regularization ป้องกัน overfitting")

    # ═══════════════════════════════════
    # 2. TRAINING DATA
    # ═══════════════════════════════════
    add_heading(doc, "2. ข้อมูลที่ใช้เทรน (Training Data)")

    add_heading(doc, "2.1 แหล่งข้อมูล", level=2)
    add_table(doc,
        ["ประเภทข้อมูล", "จำนวน", "สัดส่วน", "คำอธิบาย"],
        [
            ["Real Data (ข้อมูลจริง)", "585", "72.7%", "การตัดสินใจจาก HR 10 ตำแหน่ง 5 บริษัท"],
            ["Synthetic Data (สังเคราะห์)", "220", "27.3%", "Edge cases 12 patterns แก้ bias"],
            ["รวม", "805", "100%", ""],
        ])

    add_heading(doc, "2.2 การแบ่งข้อมูล (Train/Test Split)", level=2)
    add_table(doc,
        ["ชุดข้อมูล", "จำนวน", "สัดส่วน"],
        [
            ["Training Set", "644", "80%"],
            ["Test Set", "161", "20%"],
            ["Cross-Validation", "5-Fold Stratified", "ทั้งชุด"],
        ])
    add_para(doc, "ใช้ train_test_split ของ scikit-learn แบบ Stratified เพื่อรักษาสัดส่วน class ในแต่ละชุดข้อมูล", italic=True)

    add_heading(doc, "2.3 การกระจายข้อมูล (Data Distribution)", level=2)
    add_image(doc, "fig1_data_distribution.png", "รูปที่ 1 การกระจายข้อมูล Real Data, Synthetic Data และ Combined Data")

    add_heading(doc, "2.4 ตำแหน่งงานที่ใช้เทรน", level=2)
    add_table(doc,
        ["บริษัท", "ตำแหน่ง", "จำนวน"],
        [
            ["Digital Solutions Ltd.", "Frontend Dev, Full-Stack Dev", "113"],
            ["TechVision Solutions", "Backend Dev, Frontend Dev", "117"],
            ["DataPro Analytics", "Data Analyst, Full-Stack Dev", "119"],
            ["NetSecure Systems", "Network Engineer, Cybersecurity", "118"],
            ["AppNova Digital", "Mobile Dev, QA/Software Tester", "118"],
        ])

    # ═══════════════════════════════════
    # 3. FEATURES
    # ═══════════════════════════════════
    add_heading(doc, "3. Features ที่ใช้เทรน (17 Features)")
    add_table(doc,
        ["#", "Feature Name", "คำอธิบาย", "ประเภท"],
        [
            ["1", "skills_match_ratio", "สัดส่วนทักษะที่ตรงกับ Job Requirements", "Continuous [0,1]"],
            ["2", "skills_match_count", "จำนวนทักษะที่ตรง", "Integer"],
            ["3", "total_skills", "จำนวนทักษะทั้งหมดใน Resume", "Integer"],
            ["4", "major_match_score", "คะแนนความตรงของสาขาวิชา", "Continuous [0,1]"],
            ["5", "relevant_projects", "จำนวนโปรเจคที่เกี่ยวข้อง", "Integer"],
            ["6", "total_projects", "จำนวนโปรเจคทั้งหมด", "Integer"],
            ["7", "gpa_value", "เกรดเฉลี่ย (GPA)", "Continuous [0,4]"],
            ["8", "has_gpa", "มี GPA หรือไม่", "Binary {0,1}"],
            ["9", "has_relevant_exp", "มีประสบการณ์ที่เกี่ยวข้อง", "Binary {0,1}"],
            ["10", "has_cert", "มีใบรับรองหรือไม่", "Binary {0,1}"],
            ["11", "soft_skills_count", "จำนวน Soft Skills", "Integer"],
            ["12", "resume_completeness", "ความสมบูรณ์ของ Resume", "Continuous [0,1]"],
            ["13", "gpa_below_min", "GPA ต่ำกว่าเกณฑ์ขั้นต่ำ", "Binary {0,1}"],
            ["14", "cert_job_relevance", "ใบรับรองเกี่ยวข้องกับงาน", "Binary {0,1}"],
            ["15", "gpa_gap", "ส่วนต่าง GPA จากเกณฑ์ขั้นต่ำ", "Continuous"],
            ["16", "skill_focus_ratio", "สัดส่วน skills ที่ focus ตรงงาน", "Continuous [0,1]"],
            ["17", "soft_skills_match_ratio", "สัดส่วน soft skills ที่ตรง", "Continuous [0,1]"],
        ])

    # ═══════════════════════════════════
    # 4. HYPERPARAMETERS
    # ═══════════════════════════════════
    add_heading(doc, "4. อัลกอริทึมและ Hyperparameters")

    add_heading(doc, "4.1 อัลกอริทึม: XGBoost (XGBClassifier)", level=2)
    add_para(doc, "XGBoost (Extreme Gradient Boosting) เป็น Ensemble Learning ที่ใช้หลักการ Gradient Boosting สร้าง Decision Trees ที่ปรับปรุงต่อเนื่อง โดยแต่ละ tree จะเรียนรู้จาก error ของ tree ก่อนหน้า ช่วยให้ model มีความแม่นยำสูง")

    add_heading(doc, "4.2 Hyperparameters", level=2)
    add_table(doc,
        ["Parameter", "ค่า", "คำอธิบาย"],
        [
            ["n_estimators", "150", "จำนวนต้นไม้ (boosting rounds)"],
            ["max_depth", "4", "ความลึกสูงสุดของต้นไม้"],
            ["learning_rate", "0.08", "อัตราการเรียนรู้"],
            ["min_child_weight", "3", "น้ำหนักขั้นต่ำของ leaf node"],
            ["subsample", "0.85", "สัดส่วนข้อมูลสำหรับแต่ละต้นไม้"],
            ["colsample_bytree", "0.85", "สัดส่วน features สำหรับแต่ละต้นไม้"],
            ["reg_alpha", "0.1", "L1 regularization"],
            ["reg_lambda", "1.0", "L2 regularization"],
            ["objective", "binary:logistic", "Binary classification"],
            ["eval_metric", "logloss", "Logarithmic loss"],
            ["scale_pos_weight", "1.35", "ถ่วงน้ำหนัก class"],
            ["decision_threshold", "0.50", "เกณฑ์ตัดสินใจ"],
        ])

    # ═══════════════════════════════════
    # 5. RESULTS
    # ═══════════════════════════════════
    add_heading(doc, "5. ผลการเทรน (Training Results)")

    add_heading(doc, "5.1 Classification Metrics", level=2)
    add_table(doc,
        ["เมตริก", "ค่า", "คำอธิบาย"],
        [
            ["Accuracy", "86.3%", "ความถูกต้องโดยรวม"],
            ["Precision", "84.1%", "ความแม่นยำของการ predict Accepted"],
            ["Recall", "84.1%", "ความสามารถในการค้นหาคนที่ HR Accept จริง"],
            ["F1-Score", "84.1%", "Harmonic mean ของ Precision+Recall"],
            ["AUC-ROC", "0.954", "พื้นที่ใต้กราฟ ROC (Outstanding)"],
        ])

    add_heading(doc, "5.2 Confusion Matrix", level=2)
    add_image(doc, "fig2_confusion_matrix.png", "รูปที่ 2 Confusion Matrix แสดงผลการจำแนก")
    add_table(doc,
        ["", "Predicted Rejected", "Predicted Accepted"],
        [
            ["Actual Rejected", "TN = 81", "FP = 11"],
            ["Actual Accepted", "FN = 11", "TP = 58"],
        ])
    add_para(doc, "True Negative (81): โมเดล reject ถูกต้อง | True Positive (58): โมเดล accept ถูกต้อง")
    add_para(doc, "False Positive (11): โมเดล accept แต่ HR reject | False Negative (11): โมเดล reject แต่ HR accept")

    add_heading(doc, "5.3 ROC Curve", level=2)
    add_image(doc, "fig3_roc_curve.png", "รูปที่ 3 ROC Curve แสดง AUC = 0.954 (Outstanding Level)")
    add_para(doc, "ROC Curve แสดงความสามารถในการแยก class ที่ทุก threshold โดย AUC = 0.954 หมายความว่าโมเดลมีความสามารถในการจำแนกที่ดีมาก (เกณฑ์: >= 0.8 = Excellent, >= 0.9 = Outstanding)")

    add_heading(doc, "5.4 Precision-Recall Curve", level=2)
    add_image(doc, "fig4_precision_recall.png", "รูปที่ 4 Precision-Recall Curve")
    add_para(doc, "Precision-Recall Curve แสดง trade-off ระหว่าง Precision และ Recall โมเดลรักษา Precision สูงได้แม้ Recall สูง")

    add_heading(doc, "5.5 Feature Importance", level=2)
    add_image(doc, "fig5_feature_importance.png", "รูปที่ 5 Feature Importance แสดงตัวแปรที่มีอิทธิพลต่อการตัดสินใจ")
    add_para(doc, "Feature ที่สำคัญที่สุด: 1) Skills Match Count (0.49) 2) Skills Match Ratio (0.10) 3) Relevant Projects (0.05) ซึ่งสอดคล้องกับหลักการ HR ในการคัดกรองที่พิจารณาทักษะตรงตำแหน่งเป็นหลัก")

    add_heading(doc, "5.6 Score Distribution", level=2)
    add_image(doc, "fig10_score_distribution.png", "รูปที่ 6 Score Distribution แยกตาม class")
    add_para(doc, "กราฟ Score Distribution แสดงว่าโมเดลแยก class ได้ชัดเจน: Rejected กระจุกทางซ้าย (คะแนนต่ำ) Accepted กระจุกทางขวา (คะแนนสูง)")

    # ═══════════════════════════════════
    # 6. CROSS-VALIDATION
    # ═══════════════════════════════════
    add_heading(doc, "6. Cross-Validation (5-Fold Stratified)")
    add_image(doc, "fig7_cross_validation.png", "รูปที่ 7 ผล 5-Fold Stratified Cross-Validation")
    add_table(doc,
        ["เมตริก", "ค่า"],
        [
            ["Method", "5-Fold Stratified K-Fold"],
            ["Mean Accuracy", "87.6%"],
            ["Standard Deviation", "2.2%"],
        ])
    add_para(doc, "Cross-Validation accuracy สม่ำเสมอทุก fold (std = 2.2%) แสดงว่าโมเดลไม่ overfitting และ generalize ได้ดี")

    # ═══════════════════════════════════
    # 7. LEARNING CURVE
    # ═══════════════════════════════════
    add_heading(doc, "7. Learning Curve")
    add_image(doc, "fig6_learning_curve.png", "รูปที่ 8 Learning Curve แสดง Training vs Validation Score")
    add_para(doc, "Learning Curve แสดงว่า: Training score ลดลงเล็กน้อยเมื่อข้อมูลเพิ่ม (ไม่ overfit), Validation score เพิ่มขึ้น (ข้อมูลเพียงพอ), Gap ระหว่าง Train/Val แคบลง (generalization ดี)")

    # ═══════════════════════════════════
    # 8. COMPARISON
    # ═══════════════════════════════════
    add_heading(doc, "8. เปรียบเทียบ v4.4 vs v5.0")

    add_heading(doc, "8.1 Metrics Comparison", level=2)
    add_image(doc, "fig8_version_comparison.png", "รูปที่ 9 เปรียบเทียบ Performance ระหว่าง v4.4 และ v5.0")
    add_table(doc,
        ["เมตริก", "v4.4 (เดิม)", "v5.0 (ใหม่)", "การเปลี่ยนแปลง"],
        [
            ["Samples", "345", "805", "+133%"],
            ["Accuracy", "84.1%", "86.3%", "+2.2%"],
            ["Precision", "84.0%", "84.1%", "+0.1%"],
            ["Recall", "75.0%", "84.1%", "+9.1%"],
            ["F1-Score", "79.3%", "84.1%", "+4.8%"],
            ["AUC-ROC", "0.913", "0.954", "+0.041"],
            ["CV Mean", "86.1%", "87.6%", "+1.5%"],
        ])

    add_heading(doc, "8.2 Per-Position Accuracy", level=2)
    add_image(doc, "fig9_position_accuracy.png", "รูปที่ 10 เปรียบเทียบ Accuracy รายตำแหน่ง")
    add_table(doc,
        ["ตำแหน่ง", "v4.4", "v5.0", "เปลี่ยนแปลง"],
        [
            ["QA/Software Tester", "61.4%", "84.2%", "+22.8%"],
            ["Data Analyst", "82.5%", "94.7%", "+12.3%"],
            ["Cybersecurity", "80.7%", "93.0%", "+12.3%"],
            ["Network Engineer", "77.2%", "87.7%", "+10.5%"],
            ["Backend Developer", "75.4%", "82.5%", "+7.0%"],
            ["Full-Stack Dev (Digital)", "89.2%", "95.4%", "+6.2%"],
            ["Frontend Dev (TechVision)", "89.5%", "93.0%", "+3.5%"],
            ["Full-Stack Dev (DataPro)", "82.5%", "86.0%", "+3.5%"],
            ["Mobile Developer", "87.7%", "91.2%", "+3.5%"],
            ["Frontend Dev (Digital)", "100.0%", "100.0%", "0%"],
        ])

    add_heading(doc, "8.3 Live Test (585 Real Applications)", level=2)
    add_table(doc,
        ["เมตริก", "ค่า"],
        [
            ["Old Model Accuracy", "469/585 = 80.2%"],
            ["New v5 Accuracy", "516/585 = 88.2%"],
            ["Improvement", "+8.0%"],
            ["Fixed (wrong to right)", "57 cases"],
            ["Broken (right to wrong)", "10 cases"],
            ["Net Improvement", "+47 cases"],
        ])

    # ═══════════════════════════════════
    # 9. BIAS ANALYSIS
    # ═══════════════════════════════════
    add_heading(doc, "9. การปรับปรุง Bias")

    add_heading(doc, "9.1 ปัญหาที่พบใน v4.4", level=2)
    add_para(doc, "1. Network students ได้คะแนนสูงเกินไปในตำแหน่ง QA/Mobile/Data Analyst เพราะ GPA สูง")
    add_para(doc, "2. Software developers ถูก underrate ในตำแหน่ง QA แม้มี testing mindset")
    add_para(doc, "3. GPA สูงอย่างเดียวทำให้โมเดล accept ทั้งที่ Skills ไม่ตรง")
    add_para(doc, "4. ไม่ใส่ GPA ทำให้ถูก reject ทั้งที่ Skills ดี")

    add_heading(doc, "9.2 วิธีแก้ไข", level=2)
    add_para(doc, "สร้าง Synthetic Data 220 records จาก 12 patterns เพื่อสอนโมเดลให้เรียนรู้ edge cases ที่ real data ขาด")
    add_table(doc,
        ["Pattern", "จำนวน", "การสอน"],
        [
            ["High skills match → Accept", "25", "พื้นฐาน"],
            ["Low skills match → Reject", "25", "พื้นฐาน"],
            ["Network → Non-Network job → Reject", "30", "แก้ bias หลัก"],
            ["Multi-project dev → QA Accept", "20", "แก้ QA underrate"],
            ["Low GPA + strong skills → Accept", "15", "GPA Case 2"],
            ["Low GPA + no skills → Reject", "15", "ตรงข้าม"],
            ["Missing GPA + skills → Accept", "15", "GPA Case 1"],
            ["Missing GPA + no skills → Reject", "15", "ตรงข้าม"],
            ["High GPA + 0 skills → Reject", "20", "GPA alone != accept"],
            ["Cert + Exp → Accept", "15", "ใบรับรอง+ประสบการณ์"],
            ["Non-IT → Reject", "15", "ไม่มี IT background"],
            ["High soft skills → Accept", "10", "QA candidates"],
        ])

    add_heading(doc, "9.3 ผลลัพธ์หลังแก้ bias", level=2)
    add_table(doc,
        ["ผู้สมัคร", "ตำแหน่ง", "v4.4", "v5.0", "HR ตัดสิน"],
        [
            ["Lada (Network)", "QA Tester", "86%", "9.7%", "Reject ✓"],
            ["Piruntam (SOC)", "QA Tester", "79%", "24.9%", "Reject ✓"],
            ["Lada (Network)", "Mobile Dev", "83%", "2.9%", "Reject ✓"],
            ["Phanudach (Flutter)", "Mobile Dev", "52%", "99.6%", "Accept ✓"],
            ["GPA 3.8 + 0 skills", "ทุกตำแหน่ง", "~70%", "1.2%", "Reject ✓"],
        ])

    # ═══════════════════════════════════
    # 10. CONCLUSION
    # ═══════════════════════════════════
    add_heading(doc, "10. สรุป (Conclusion)")
    add_para(doc, "XGBoost v5.0 สำหรับระบบคัดกรอง Resume อัตโนมัติได้รับการปรับปรุงอย่างมีนัยสำคัญ:")
    add_para(doc, "• Accuracy เพิ่มขึ้นจาก 84.1% เป็น 86.3% (+2.2%)")
    add_para(doc, "• Recall เพิ่มขึ้นมากจาก 75.0% เป็น 84.1% (+9.1%)")
    add_para(doc, "• AUC-ROC ถึง 0.954 (Outstanding Level)")
    add_para(doc, "• Live test accuracy ดีขึ้นจาก 80.2% เป็น 88.2% (+8.0%)")
    add_para(doc, "• แก้ bias ของ cross-domain scoring สำเร็จ (57 cases fixed, net +47)")
    add_para(doc, "• โมเดลไม่ overfitting (CV std = 2.2%)")
    add_para(doc, "• รองรับ 10 ตำแหน่งใน 5 บริษัทได้อย่างมีประสิทธิภาพ")

    # Save
    doc.save(str(OUTPUT_PATH))
    print(f"✅ DOCX saved: {OUTPUT_PATH}")
    print(f"   Size: {OUTPUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
